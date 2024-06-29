from typing import List
from fastapi import Response, status, HTTPException, WebSocket, APIRouter, Depends, Body, WebSocketDisconnect
from telegram_sender import TelegramSender
from broker_manager import BrokerManager
from deps import get_current_user_impl, get_current_user
import json

from database_manager import DatabaseManager
from docx import Document
import websockets
import json
import io
from schemas import *
from bestconfig import Config


router = APIRouter()

config = Config()

queue_name = 'apiparser'
broker_host = config['BROKER_HOST']

@router.get("/download/all/")
async def download_all_posts():
    try:
        # Преобразование данных в .doc
        doc = Document()
        doc.add_heading('Сгенерированные посты', level=1)

        databaseManager = DatabaseManager(config['DB_CONNECTION'])
        posts = databaseManager.get_all_posts()
        for post in posts:
            doc.add_heading(post['title'], level=2)
            doc.add_paragraph(post['article'])
            doc.add_paragraph(post['url'])
            doc.add_paragraph('')
            doc.add_paragraph('')
        
        # Создание байтового потока для хранения .doc файла
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Отправка файла в ответе
        return Response(doc_stream.getvalue(), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": "attachment;filename=posts.docx"})
    except Exception as e:
        print(f"Error generating or sending .doc file: {e}")
        return Response(status_code=500)


@router.get("/all/")
async def read_all():
    databaseManager = DatabaseManager(config['DB_CONNECTION'])
    result = databaseManager.get_all_posts()
    return result


@router.get("/all/del/")
async def delete_all():
    databaseManager = DatabaseManager(config['DB_CONNECTION'])
    result = databaseManager.delete_all_posts()
    return result


@router.post("/del/")
async def delete_item(item: ItemNumber):
    databaseManager = DatabaseManager(config['DB_CONNECTION'])
    result = databaseManager.delete_post_by_id(item.number)
    return result


@router.post("/create/")
async def create_item(item: Item):
    databaseManager = DatabaseManager(config['DB_CONNECTION'])
    result = databaseManager.create_post(item.name)
    return result


@router.post("/run")
async def create_item(item: Item):
    broker = BrokerManager(queue_name, broker_host)
    msg = json.dumps({'type': 'standart','resource': item.name})
    broker.send_msg(msg)
    broker.close()
    return item.name


# DONE
# websockets
connected_websockets = {}

async def send_message(user_id:int, message: str, message_type: str):
    if str(user_id) in connected_websockets:
        msg = json.dumps({'message': message, 'type': message_type})
        ws_connection = connected_websockets[str(user_id)]
        await ws_connection.send_text(msg)

@router.websocket("/posts/progress/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            data = await websocket.receive_text()
            print("recived data:", data)
            data = json.loads(data)
            print("testtest: ", (data['token']))

            # getting user email
            user = get_current_user(data['token'])
            if user is None:
                return
            
            # saving in dict
            connected_websockets[str(user.id)] = websocket

            # await send_message(data)
    except websockets.exceptions.ConnectionClosedError:
        for key, value in dict(connected_websockets).items():
            if value == websocket:
                del connected_websockets[key]
    except WebSocketDisconnect:
        for key, value in dict(connected_websockets).items():
            if value == websocket:
                del connected_websockets[key]
################

# DONE
@router.get("/posts/titles")
async def get_titles_from_db():
    print("geting titles from db...")
    db = DatabaseManager(config['DB_CONNECTION'])
    posts = db.get_all_posts()
    print('sending (', len(posts), ' items)...')
    return posts


# DONE
@router.post("/posts/titles/test")
async def parse_titles(request: ParseTitlesRequest, user: SystemUser = Depends(get_current_user_impl)):
    parsers = {
        'SCIENTIFICAMERICAN': 'https://www.scientificamerican.com/artificial-intelligence/',
        'MIT': 'https://news.mit.edu/topic/artificial-intelligence2',
        'EXTREMETECH': 'https://www.extremetech.com/tag/artificial-intelligence',
        'VENTUREBEAT': 'https://venturebeat.com/category/ai/',
        'GIZMODO': 'https://gizmodo.com/moderna-ceo-chatgpt-employees-vaccines-openai-1851435620',
        'SYNCED': 'https://syncedreview.com/category/popular/',
    }

    processed_urls = []
    for item in request.resources:
        if item == 'all':
            processed_urls = list(parsers.values())
            break
                
        if item in parsers:
            processed_urls.append(parsers[item])

    for item in request.urls:
        processed_urls.append(item)

    msg = json.dumps({'type': 'titles',
                      'resources': processed_urls, 
                      'period_days': request.period_days,
                      'user_id': user.id,
                    })
    
    broker = BrokerManager(queue_name, broker_host)
    broker.send_msg(msg)
    broker.close()
    print("отправленно!!!!!!!!")
    return {"message": "Запрос успешно обработан"}


@router.post("/posts/sendprogress")
async def send_progress(request: ProgressRequest):
    print('sended progress')
    await send_message(request.user_id, request.status, 'progress')
    print('sended!')
    

# articles

@router.post("/posts/article")
async def parse_titles(request: ParseArticleRequest, user: SystemUser = Depends(get_current_user_impl)):
    msg = json.dumps({'type': 'articles',
                      'resources': request.url,
                      'user_id': user.id
                    })
    broker = BrokerManager(queue_name, broker_host)
    broker.send_msg(msg)
    broker.close()
    print("отправленно!!!!!!!!")
    return {"message": "Запрос успешно обработан"}


@router.post("/posts/sendpost")
async def send_progress(request: PostRequest):
    await send_message(request.user_id, request.content, 'post')


# answers and posts

@router.post("/posts/telegram")
async def create_item(request: PostSchema):
    db = DatabaseManager(config['DB_CONNECTION'])
    tg_sender = TelegramSender()
    message_id_list = await tg_sender.send_message(post=request.post)
    for message_id in message_id_list:
        db.add_message_id(post_id=request.post_id, message_id=message_id)
    return

@router.get("/answers/all")
async def get_all_answers():
    db = DatabaseManager(config['DB_CONNECTION'])
    result = db.get_all_answers()
    print('****************************')
    print('****************************')
    print('****************************')
    print(result)
    if result is None:
        return []
    return result
